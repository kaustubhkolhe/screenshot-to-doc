import os
from docx import Document
from docx.shared import Inches
import pyautogui
from pynput.keyboard import Key, Listener
import datetime

current_dir = os.getcwd()
document = Document()  # Initialize document here
screenshot_shortcut = Key.print_screen
exit_combination_msg = 'left ctrl + space'
exit_combination = {Key.ctrl_l, Key.space}
default_doc_name = "testingScreenShots"
img_count = 1
master_path = ""
file_name = ""
currently_pressed = set()
start_time = None
end_time = None  # Initialize end_time variable


def on_press(key):
    global start_time
    check_key(key)
    if key in exit_combination:
        currently_pressed.add(key)
        if currently_pressed == exit_combination:
            # Capture end time when exit combination is pressed
            capture_end_time()
            listener.stop()
            exit_fun()


def on_release(key):
    try:
        currently_pressed.remove(key)
    except KeyError:
        pass


def check_key(key):
    if key == screenshot_shortcut:
        save_img()


def add_ss_to_doc():
    global document
    img_path = os.path.join(master_path, "shots", f"{img_count}.png")
    document.add_paragraph().add_run().add_picture(img_path, width=Inches(5.90551), height=Inches(3.54331))
    document.add_paragraph(f"Screenshot {img_count}")


def save_doc(doc_file_name):
    global document
    document.save(os.path.join(master_path, f"{doc_file_name}.docx"))
    print("Document saved at:")
    print(os.path.join(master_path, f"{doc_file_name}.docx"))


def exit_fun():
    global file_name
    print("<===================================>")
    if not file_name:
        print(f"Saving Document with default name ({default_doc_name})")
        file_name = default_doc_name

    save_doc(file_name)


def get_directory_name():
    global master_path, file_name, start_time
    while True:
        directory_name = input("Enter directory name: ")
        master_path = os.path.join(current_dir, "TestingData", directory_name)
        if os.path.exists(master_path):
            print("Folder already exists. Please choose a different name.")
        else:
            os.makedirs(os.path.join(master_path, "shots"))
            break

    file_name = input("Enter document name: ")
    print("Press PrtSc to take the Screenshot and save to folder and To Document")
    print(f"Press {exit_combination_msg} to exit and save the document")
    print("_______________________")
    print("Current Directory is =>", os.getcwd())


def save_img():
    global img_count
    shot = pyautogui.screenshot()
    path = os.path.join(master_path, "shots")
    try:
        shot.save(os.path.join(path, f"{img_count}.png"))
        print(f'File Saved as {path}\\{img_count}.png')
        add_ss_to_doc()
        img_count += 1
    except Exception as e:
        print(f'Error occurred while saving screenshot: {e}')


def print_start_msg():
    global start_time

    title = input("Enter document title: ")
    start_time = datetime.datetime.now()
    section = document.sections[0]
    header = section.header
    header.paragraphs[0].text = title
    header.add_paragraph(f"Start Time: {start_time.strftime('%Y-%m-%d %H:%M:%S')}")

    get_directory_name()


def capture_end_time():
    global end_time, img_count

    end_time = datetime.datetime.now()
    section = document.sections[0]
    header = section.header
    header.add_paragraph(f"End Time: {end_time.strftime('%Y-%m-%d %H:%M:%S')}")
    header.add_paragraph(f"Total Screenshots: {img_count - 1}")


# Collect events until released
with Listener(on_press=on_press, on_release=on_release) as listener:
    print_start_msg()
    listener.join()
