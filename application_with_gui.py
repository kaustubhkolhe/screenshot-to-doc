import os
import datetime
import pyautogui
from docx import Document
from docx.shared import Inches
from pynput.keyboard import Key, Listener
import tkinter as tk
from tkinter import messagebox
import threading
import psutil

process_started = False
current_dir = os.getcwd()
document = Document()
screenshot_shortcut = Key.print_screen
img_count = 1
master_path = ""
file_name = ""
currently_pressed = set()
start_time = None
end_time = None
title = None
listener = None


# Function to print start message and initiate the process
def print_start_msg(title_value, directory_name, doc_name):
    global start_time, title
    print("Document Title:", title_value)
    title = title_value
    start_time = datetime.datetime.now()
    get_directory_name(directory_name, doc_name)


# Callback function for key press events
def on_press(key):
    global currently_pressed

    currently_pressed = set()
    check_key(key)


# Function to check if the pressed key is the screenshot shortcut
def check_key(key):
    if key == screenshot_shortcut:
        save_img()


# Function to add screenshot to the document
def add_ss_to_doc(img_path):
    document.add_paragraph().add_run().add_picture(img_path, width=Inches(5.90551), height=Inches(3.54331))


# Function to save the document
def save_doc(doc_file_name):
    try:
        document.save(os.path.join(master_path, f"{doc_file_name}.docx"))
        print("Document saved at:")
        print(os.path.join(master_path, f"{doc_file_name}.docx"))
    except Exception as e:
        print(f"Error saving document: {e}")


# Function to handle program exit
def exit_fun():
    global file_name
    print("<===================================>")

    save_doc(file_name)


# Function to get the directory name
def get_directory_name(directory_name, doc_name):
    global master_path, file_name

    master_path = os.path.join(current_dir, "TestingData", directory_name)

    if not os.path.exists(master_path):
        os.makedirs(master_path)
        os.makedirs(os.path.join(master_path, "Screenshots"))
    else:
        print("Directory already exists. Using existing directory.")
        if not os.path.exists(os.path.join(master_path, "Screenshots")):
            os.makedirs(os.path.join(master_path, "Screenshots"))

    file_name = doc_name
    print("Press PrtSc to take the Screenshot and save to folder and To Document")

    print("_______________________")
    print("Current Directory is =>", os.getcwd())


# Function to save the screenshot
def save_img():
    global img_count
    shot = pyautogui.screenshot()
    path = os.path.join(master_path, "Screenshots")
    try:
        document_folder = os.path.join(path, file_name)
        if not os.path.exists(document_folder):
            os.makedirs(document_folder)

        screenshot_path = os.path.join(document_folder, f"{img_count}.png")
        shot.save(screenshot_path)
        print(f'File Saved as {screenshot_path}')
        add_ss_to_doc(screenshot_path)
        img_count += 1
    except Exception as e:
        print(f'Error occurred while saving screenshot: {e}')


# Function to capture end time and finalize the document
def capture_end_time():
    global end_time, img_count, document, title, start_time
    end_time = datetime.datetime.now()

    if document.paragraphs:
        first_paragraph = document.paragraphs[0]
    else:
        first_paragraph = document.add_paragraph()

    total_screenshots_paragraph = first_paragraph.insert_paragraph_before(
        "_________________________________________________________________________________________________________")
    total_screenshots_paragraph.style.paragraph_format.line_spacing = 1

    title_paragraph = first_paragraph.insert_paragraph_before()
    title_paragraph.add_run("Document Title: ")
    title_paragraph.add_run(title)
    document.paragraphs.insert(0, title_paragraph)

    total_screenshots_paragraph = first_paragraph.insert_paragraph_before(f"Total Screenshots: {img_count - 1}")
    total_screenshots_paragraph.style.paragraph_format.line_spacing = 1

    start_time_str = start_time.strftime('%Y-%m-%d %H:%M:%S') if start_time else "Not available"
    end_time_str = end_time.strftime('%Y-%m-%d %H:%M:%S') if end_time else "Not available"

    start_time_paragraph = first_paragraph.insert_paragraph_before(
        f"Start Time: {start_time_str}")
    start_time_paragraph.style.paragraph_format.line_spacing = 1

    end_time_paragraph = first_paragraph.insert_paragraph_before(f"End Time: {end_time_str}")
    end_time_paragraph.style.paragraph_format.line_spacing = 1

    total_screenshots_paragraph = first_paragraph.insert_paragraph_before(
        "_________________________________________________________________________________________________________\n"
        "\n\n")
    total_screenshots_paragraph.style.paragraph_format.line_spacing = 1


# Function to terminate the program
def terminate_program():
    global listener, document, process_started, master_path, file_name, start_time, end_time, title, img_count

    try:
        if listener:
            listener.stop()  # Stop the listener if it exists
            listener.join()  # Wait for the listener thread to terminate

        try:
            capture_end_time()
            exit_fun()
        except Exception as e:
            print(f"Error occurred during termination: {e}")

        # Reset all global variables
        document = Document()  # Reset document to default
        process_started = False  # Reset process_started to default
        master_path = ""  # Reset master_path to default
        file_name = ""  # Reset file_name to default
        start_time = None  # Reset start_time to default
        end_time = None  # Reset end_time to default
        title = None  # Reset title to default
        img_count = 1  # Reset img_count to default

        exit()

    except Exception as e:
        print(f"Error occurred during termination: {e}")
        exit()


# Function to start the listener
def start_listener():
    global listener
    listener = Listener(on_press=on_press)
    listener.start()
    listener.join()


# Front end started
def start_process():
    global process_started, file_name, start_button
    document_title = doc_title_entry.get()
    directory_name = dir_name_entry.get().strip()
    file_name = file_name_entry.get().strip()

    if not validate_directory_name(directory_name):
        messagebox.showerror("Error", "Illegal character in directory name")
        return
    if not validate_file_name(file_name):
        messagebox.showerror("Error", "Illegal character in file name")
        return

    print("Document Title:", document_title)
    print("Directory Name:", directory_name)
    print("File Name:", file_name)
    print_start_msg(document_title, directory_name, file_name)

    # Start the listener on a separate thread
    listener_thread = threading.Thread(target=start_listener)
    listener_thread.start()

    start_button.config(state=tk.DISABLED)
    stop_button.config(state=tk.NORMAL)
    process_started = True
    doc_title_entry.config(state=tk.DISABLED)
    dir_name_entry.config(state=tk.DISABLED)
    file_name_entry.config(state=tk.DISABLED)
    clear_button.config(state=tk.DISABLED)


# Function to stop the process
def stop_process():
    global process_started, start_button
    print("Process Stopped")

    # Start the termination process on a separate thread
    stop_thread = threading.Thread(target=terminate_program)
    stop_thread.start()

    start_button.config(state=tk.NORMAL)
    stop_button.config(state=tk.DISABLED)
    process_started = False
    doc_title_entry.config(state=tk.NORMAL)
    dir_name_entry.config(state=tk.NORMAL)
    file_name_entry.config(state=tk.NORMAL)
    clear_button.config(state=tk.NORMAL)


# Function to validate directory name
def validate_directory_name(directory_name):
    invalid_chars = set('/\\?%*:|"<>')
    return not any(char in invalid_chars for char in directory_name)


# Function to validate file name
def validate_file_name(file_name):
    invalid_chars = set('/\\?%*:|"<>')
    return not any(char in invalid_chars for char in file_name)


# Function to clear entries in GUI
def clear_entries():
    doc_title_entry.delete(0, 'end')
    dir_name_entry.delete(0, 'end')
    file_name_entry.delete(0, 'end')
    update_start_button_state()


# Function to update the state of the start button
def update_start_button_state(event=None):
    global process_started, start_button
    if all((doc_title_entry.get(), dir_name_entry.get(), file_name_entry.get())) and not process_started:
        start_button.config(state=tk.NORMAL)
    else:
        start_button.config(state=tk.DISABLED)


# Function to check and start the process
def check_and_start(event=None):
    if all((doc_title_entry.get(), dir_name_entry.get(), file_name_entry.get())) and not process_started:
        start_process()


# Function to handle closing event of the application
def on_closing():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):

        try:
            psutil.Process(os.getpid()).terminate()  # Terminate the current process
        except Exception as e:
            print(f"Error occurred during process termination: {e}")
        root.destroy()  # Destroy the Tkinter window


# GUI setup
root = tk.Tk()
root.title("Capture Screenshots")  # Display application name
root.geometry("450x350")
root.resizable(False, False)

ENTRY_FONT = ("System", 14)
NOTE_FONT = ("System", 10)
BUTTON_FONT = ("Lato", 14)
TEXT_FONT = ("Courier", 13)

# Labels and Entries
doc_title_label = tk.Label(root, text="Document Title:", font=ENTRY_FONT)
doc_title_label.grid(row=0, column=0, pady=(20, 5), sticky='w')
doc_title_entry = tk.Entry(root, font=TEXT_FONT)
doc_title_entry.grid(row=0, column=1, pady=(20, 5), padx=(0, 20), sticky='ew')

dir_name_label = tk.Label(root, text="Directory Name:", font=ENTRY_FONT)
dir_name_label.grid(row=1, column=0, pady=5, sticky='w')
dir_name_entry = tk.Entry(root, font=TEXT_FONT)
dir_name_entry.grid(row=1, column=1, pady=5, padx=(0, 20), sticky='ew')

file_name_label = tk.Label(root, text="File Name:", font=ENTRY_FONT)
file_name_label.grid(row=2, column=0, pady=(5, 20), sticky='w')
file_name_entry = tk.Entry(root, font=TEXT_FONT)
file_name_entry.grid(row=2, column=1, pady=(5, 20), padx=(0, 20), sticky='ew')

# Buttons
clear_button = tk.Button(root, text="Clear", command=clear_entries, font=BUTTON_FONT)
clear_button.grid(row=3, column=0, pady=5, padx=(20, 10), sticky='ew')

start_button = tk.Button(root, text="Start", command=check_and_start, font=BUTTON_FONT, state=tk.DISABLED)
start_button.grid(row=3, column=1, pady=5, padx=(10, 20), sticky='ew')

stop_button = tk.Button(root, text="Stop", command=stop_process, font=BUTTON_FONT, state=tk.DISABLED)
stop_button.grid(row=4, column=1, pady=5, padx=(10, 20), sticky='ew')

# Event bindings
doc_title_entry.bind("<KeyRelease>", update_start_button_state)
dir_name_entry.bind("<KeyRelease>", update_start_button_state)
file_name_entry.bind("<KeyRelease>", update_start_button_state)

# Alternate way to start by pressing enter key
root.bind("<Return>", check_and_start)

# Procedure note
procedure_label = tk.Label(root,
                           text="After clicking start, begin taking screenshots.\npress stop to save in a word document.",
                           font=NOTE_FONT, fg="grey")
procedure_label.grid(row=5, column=0, columnspan=2, pady=(20, 0))

root.protocol("WM_DELETE_WINDOW", on_closing)

root.mainloop()
